"""
Document Parser Service
Handles ingestion and text extraction from PDF and DOCX files.
"""

import os
import PyPDF2
import docx
from typing import Tuple


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file."""
    text_parts = []
    try:
        reader = PyPDF2.PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file."""
    text_parts = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    return "\n\n".join(text_parts)


def parse_document(file_path: str) -> Tuple[str, str]:
    """
    Parse a document and return its text content and detected type.
    Returns: (text_content, file_type)
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        return text, "pdf"
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(file_path)
        return text, "docx"
    elif ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return text, "text"
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def chunk_text(text: str, max_chunk_size: int = 4000, overlap: int = 200) -> list:
    """
    Split text into overlapping chunks for LLM processing.
    Useful for large documents that exceed context window limits.
    """
    if len(text) <= max_chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chunk_size
        # Try to break at a paragraph or sentence boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind("\n\n", start, end)
            if para_break > start + max_chunk_size // 2:
                end = para_break + 2
            else:
                # Look for sentence break
                sent_break = text.rfind(". ", start, end)
                if sent_break > start + max_chunk_size // 2:
                    end = sent_break + 2

        chunks.append(text[start:end])
        start = end - overlap

    return chunks
